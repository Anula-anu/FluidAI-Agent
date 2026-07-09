from docx import Document
from docx.shared import Pt
import os

def create_word(document_text, filename="output.docx"):

    os.makedirs("docs", exist_ok=True)

    path = os.path.join("docs", filename)

    doc = Document()

    title = doc.add_heading("AI Generated Business Document", level=1)
    title.runs[0].font.size = Pt(18)

    for line in document_text.split("\n"):

        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)

        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)

        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)

        else:
            doc.add_paragraph(line)

    doc.save(path)

    return path